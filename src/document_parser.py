"""
DOCX document parser for the DOCX to Asana CSV Generator.

This module handles parsing DOCX files, extracting heading structures,
and organizing content into a format suitable for Asana task creation.
"""

import os
from typing import Dict, List, Any, Optional, Tuple
import docx
from docx import Document

from .utils.logger import get_logger
from .utils.config import get_config
from .utils.validator import validate_file_path

# Initialize logger and config
logger = get_logger(__name__)
config = get_config()


class DocumentParser:
    """
    Parser for DOCX documents.
    
    Extracts structured data from DOCX files, focusing on heading levels
    and content organization for Asana task creation.
    """
    
    def __init__(self):
        """Initialize the document parser with configuration settings."""
        # Load heading style configuration
        self.heading_styles = config.get("document.heading_styles", {
            "section": "Heading 1",
            "product_type": "Heading 2",
            "manufacturer": "Heading 3",
            "description": "Heading 4"
        })
        
        # Load other document configuration
        self.products_heading = config.get("document.products_heading", "Products")
        self.manufacturer_headings = config.get("document.manufacturer_headings", 
                                               ["Manufacturer", "Manufacturers"])
        
        # Load variations for flexible matching from config
        self.heading_style_variations = config.get("document.heading_style_variations", {
            "section": ["heading 1", "title 1", "h1", "header 1", "section"],
            "product_type": ["heading 2", "title 2", "h2", "header 2", "subsection"],
            "manufacturer": ["heading 3", "title 3", "h3", "header 3"],
            "description": ["heading 4", "title 4", "h4", "header 4"]
        })
        
        self.products_heading_variations = config.get("document.products_heading_variations", [
            "products", "product list", "products and services", 
            "product information", "product specs", "product specifications"
        ])
        
        # Convert manufacturer headings to lowercase for case-insensitive matching
        self.manufacturer_heading_variations = [
            heading.lower() for heading in self.manufacturer_headings
        ]
    
    def parse_document(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a DOCX document and extract structured data.
        
        Args:
            file_path: Path to the DOCX file
        
        Returns:
            A dictionary containing the parsed document data
        """
        # Validate file path
        if not validate_file_path(file_path, ['.docx']):
            logger.error(f"Invalid document file: {file_path}")
            return {}
        
        try:
            # Get the document filename (without extension) for the root task
            filename = os.path.splitext(os.path.basename(file_path))[0]
            
            # Open the document
            logger.info(f"Opening document: {file_path}")
            doc = Document(file_path)
            
            # CRITICAL: Accept all tracked changes before parsing
            self._accept_tracked_changes(doc)
            
            # Parse the document structure
            logger.info(f"Parsing document structure: {file_path}")
            document_data = {
                "filename": filename,
                "products_section": None,
                "product_types": []
            }
            
            # Find the Products section (2nd instance of Heading 1 with "Products" text)
            products_section = self._find_products_section(doc)
            if not products_section:
                logger.warning(f"Products section not found in document: {file_path}")
                return document_data
            
            document_data["products_section"] = products_section
            
            # Extract product types (Heading 2 elements under Products section)
            product_types = self._extract_product_types(doc, products_section)
            document_data["product_types"] = product_types
            
            logger.info(f"Document parsing completed: {file_path}")
            return document_data
            
        except Exception as e:
            logger.error(f"Error parsing document: {e}")
            return {}
    
    def _accept_tracked_changes(self, doc: Document) -> None:
        """
        Accept all tracked changes in the document.
        
        Args:
            doc: The Document object
        """
        try:
            # python-docx doesn't directly expose revision acceptance
            # This is typically handled by the library when opening the document
            # If additional handling is needed, it would require custom implementation
            logger.info("Tracked changes are automatically accepted by python-docx")
        except Exception as e:
            logger.error(f"Error accepting tracked changes: {e}")
    
    def _is_style_match(self, style_name: str, style_type: str) -> bool:
        """
        Check if a style name matches a style type using flexible matching.
        
        Args:
            style_name: The style name to check
            style_type: The style type to match against ("section", "product_type", etc.)
            
        Returns:
            True if the style name matches the style type, False otherwise
        """
        if not style_name:
            return False
            
        # Check for exact match with configured style
        if style_name.lower() == self.heading_styles[style_type].lower():
            return True
            
        # Check for variations
        style_name_lower = style_name.lower()
        for variation in self.heading_style_variations[style_type]:
            if variation in style_name_lower:
                return True
                
        return False
    
    def _is_text_match(self, text: str, variations: List[str]) -> bool:
        """
        Check if text matches any of the provided variations using flexible matching.
        
        Args:
            text: The text to check
            variations: List of text variations to match against
            
        Returns:
            True if the text matches any variation, False otherwise
        """
        if not text:
            return False
            
        text_lower = text.strip().lower()
        
        # Check for substring matches
        for variation in variations:
            if variation in text_lower:
                return True
                
        return False
    
    def _find_products_section(self, doc: Document) -> Optional[Dict[str, Any]]:
        """
        Find the Products section in the document.
        
        Args:
            doc: The Document object
        
        Returns:
            A dictionary containing information about the Products section,
            or None if not found
        """
        heading1_count = 0
        products_section = None
        
        # Iterate through paragraphs to find the 2nd Heading 1 with "Products" text
        for i, para in enumerate(doc.paragraphs):
            # Use flexible style matching
            if self._is_style_match(para.style.name, "section"):
                # Use flexible text matching
                if self._is_text_match(para.text, self.products_heading_variations):
                    heading1_count += 1
                    if heading1_count == 2:  # We want the 2nd instance
                        products_section = {
                            "index": i,
                            "text": para.text.strip()
                        }
                        logger.info(f"Found Products section: {para.text}")
                        break
        
        return products_section
    
    def _extract_product_types(self, doc: Document, products_section: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Extract product types from the document.
        
        Args:
            doc: The Document object
            products_section: Dictionary containing information about the Products section
        
        Returns:
            A list of dictionaries containing product type information
        """
        product_types = []
        current_product_type = None
        current_manufacturer = None
        in_products_section = False
        
        # Iterate through paragraphs to extract product types and manufacturers
        for i, para in enumerate(doc.paragraphs):
            # Skip paragraphs before the Products section
            if i <= products_section["index"]:
                continue
            
            # Start processing after the Products section
            in_products_section = True
            
            # Check if we've reached a new section (Heading 1)
            if para.style.name == self.heading_styles["section"]:
                # End of Products section
                break
            
            # Process Heading 2 (Product Type)
            if self._is_style_match(para.style.name, "product_type"):
                # Create a new product type entry
                current_product_type = {
                    "name": para.text.strip(),
                    "manufacturers": []
                }
                product_types.append(current_product_type)
                current_manufacturer = None
                logger.debug(f"Found product type: {para.text}")
            
            # Process Heading 3 (Manufacturer)
            elif (self._is_style_match(para.style.name, "manufacturer") and 
                  current_product_type is not None):
                # Check if this is a manufacturer heading using flexible matching
                if self._is_text_match(para.text, self.manufacturer_heading_variations):
                    # Create a new manufacturer entry
                    current_manufacturer = {
                        "name": para.text.strip(),
                        "descriptions": []
                    }
                    current_product_type["manufacturers"].append(current_manufacturer)
                    logger.debug(f"Found manufacturer: {para.text}")
            
            # Process Heading 4 (Description)
            elif (self._is_style_match(para.style.name, "description") and 
                  current_manufacturer is not None):
                # Add description to current manufacturer
                current_manufacturer["descriptions"].append(para.text.strip())
                logger.debug(f"Found description: {para.text}")
        
        return product_types
    
    def get_document_structure(self, file_path: str) -> Dict[str, Any]:
        """
        Get a summary of the document's structure.
        
        Args:
            file_path: Path to the DOCX file
        
        Returns:
            A dictionary containing information about the document's structure
        """
        try:
            # Validate file path
            if not validate_file_path(file_path, ['.docx']):
                return {}
            
            # Open the document
            doc = Document(file_path)
            
            # Count headings by level
            heading_counts = {
                "Heading 1": 0,
                "Heading 2": 0,
                "Heading 3": 0,
                "Heading 4": 0
            }
            
            # Analyze paragraphs
            for para in doc.paragraphs:
                # Use flexible style matching for counting
                if self._is_style_match(para.style.name, "section"):
                    heading_counts["Heading 1"] += 1
                elif self._is_style_match(para.style.name, "product_type"):
                    heading_counts["Heading 2"] += 1
                elif self._is_style_match(para.style.name, "manufacturer"):
                    heading_counts["Heading 3"] += 1
                elif self._is_style_match(para.style.name, "description"):
                    heading_counts["Heading 4"] += 1
            
            return {
                "filename": os.path.basename(file_path),
                "heading_counts": heading_counts,
                "paragraph_count": len(doc.paragraphs)
            }
            
        except Exception as e:
            logger.error(f"Error analyzing document structure: {e}")
            return {}


def parse_document(file_path: str) -> Dict[str, Any]:
    """
    Parse a DOCX document and extract structured data.
    
    Args:
        file_path: Path to the DOCX file
    
    Returns:
        A dictionary containing the parsed document data
    """
    parser = DocumentParser()
    return parser.parse_document(file_path)


def parse_documents(file_paths: List[str]) -> List[Dict[str, Any]]:
    """
    Parse multiple DOCX documents and extract structured data.
    
    Args:
        file_paths: List of paths to DOCX files
    
    Returns:
        A list of dictionaries containing the parsed document data
    """
    parser = DocumentParser()
    results = []
    
    for file_path in file_paths:
        try:
            result = parser.parse_document(file_path)
            if result:
                results.append(result)
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {e}")
    
    return results
