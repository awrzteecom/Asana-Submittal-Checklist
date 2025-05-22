"""
Tests for the document parser module.

This module contains tests for the document parsing functionality,
focusing on heading extraction and document structure analysis.
"""

import os
import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from src.document_parser import DocumentParser, parse_document
from src.utils.config import get_config

# Test constants
TEST_DOCS_DIR = os.path.join(os.path.dirname(__file__), "sample_documents")


def create_test_document(filename, content_structure):
    """
    Create a test DOCX document with the specified structure.
    
    Args:
        filename: Name of the document file
        content_structure: List of tuples (text, style_name) for paragraphs
    
    Returns:
        Path to the created document
    """
    # Create a new document
    doc = Document()
    
    # Add styles if they don't exist
    styles = doc.styles
    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        if style_name not in styles:
            styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    
    # Add paragraphs with specified styles
    for text, style_name in content_structure:
        p = doc.add_paragraph(text)
        p.style = doc.styles[style_name]
    
    # Save the document
    os.makedirs(TEST_DOCS_DIR, exist_ok=True)
    file_path = os.path.join(TEST_DOCS_DIR, filename)
    doc.save(file_path)
    
    return file_path


def test_document_parser_initialization():
    """Test that the DocumentParser initializes correctly."""
    parser = DocumentParser()
    
    # Check that the parser loaded configuration
    assert parser.heading_styles is not None
    assert parser.products_heading == "Products"
    assert isinstance(parser.manufacturer_headings, list)
    assert len(parser.manufacturer_headings) > 0


def test_find_products_section():
    """Test finding the Products section in a document."""
    # Create a test document with two Products sections
    content_structure = [
        ("Introduction", "Heading 1"),
        ("This is an introduction", "Normal"),
        ("Products", "Heading 1"),  # First Products section
        ("First product", "Heading 2"),
        ("Details", "Normal"),
        ("Another Section", "Heading 1"),
        ("Products", "Heading 1"),  # Second Products section (this is the one we want)
        ("Product Type 1", "Heading 2"),
        ("Manufacturer: Company A", "Heading 3"),
        ("Description 1", "Heading 4"),
    ]
    
    file_path = create_test_document("test_products_section.docx", content_structure)
    
    # Parse the document
    parser = DocumentParser()
    doc = Document(file_path)
    products_section = parser._find_products_section(doc)
    
    # Check that the second Products section was found
    assert products_section is not None
    assert products_section["text"] == "Products"
    assert products_section["index"] == 6  # 0-based index of the second Products heading

def test_find_products_section_with_variations():
    """Test finding the Products section with different text variations."""
    # Create a test document with two Products sections using different text
    content_structure = [
        ("Introduction", "Heading 1"),
        ("This is an introduction", "Normal"),
        ("Product List", "Heading 1"),  # First Products section with variation
        ("First product", "Heading 2"),
        ("Details", "Normal"),
        ("Another Section", "Heading 1"),
        ("Products and Services", "Heading 1"),  # Second Products section with variation
        ("Product Type 1", "Heading 2"),
        ("Manufacturer: Company A", "Heading 3"),
        ("Description 1", "Heading 4"),
    ]
    
    file_path = create_test_document("test_products_variations.docx", content_structure)
    
    # Parse the document
    parser = DocumentParser()
    doc = Document(file_path)
    products_section = parser._find_products_section(doc)
    
    # Check that the second Products section was found
    assert products_section is not None
    assert products_section["text"] == "Products and Services"
    assert products_section["index"] == 6  # 0-based index of the second Products heading

def test_style_matching():
    """Test flexible style matching."""
    parser = DocumentParser()
    
    # Test section style matching
    assert parser._is_style_match("Heading 1", "section") is True
    assert parser._is_style_match("heading 1", "section") is True
    assert parser._is_style_match("Title 1", "section") is True
    assert parser._is_style_match("H1", "section") is True
    assert parser._is_style_match("Custom Section Style", "section") is True
    assert parser._is_style_match("Normal", "section") is False
    
    # Test product_type style matching
    assert parser._is_style_match("Heading 2", "product_type") is True
    assert parser._is_style_match("heading 2", "product_type") is True
    assert parser._is_style_match("Title 2", "product_type") is True
    assert parser._is_style_match("H2", "product_type") is True
    assert parser._is_style_match("Custom Subsection Style", "product_type") is True
    assert parser._is_style_match("Normal", "product_type") is False

def test_text_matching():
    """Test flexible text matching."""
    parser = DocumentParser()
    
    # Test products heading matching
    assert parser._is_text_match("Products", parser.products_heading_variations) is True
    assert parser._is_text_match("Product List", parser.products_heading_variations) is True
    assert parser._is_text_match("Products and Services", parser.products_heading_variations) is True
    assert parser._is_text_match("Product Information", parser.products_heading_variations) is True
    assert parser._is_text_match("Random Text", parser.products_heading_variations) is False
    
    # Test manufacturer heading matching
    assert parser._is_text_match("Manufacturer: ABC Corp", parser.manufacturer_heading_variations) is True
    assert parser._is_text_match("Manufacturers: XYZ Inc.", parser.manufacturer_heading_variations) is True
    assert parser._is_text_match("Mfg: Company A", parser.manufacturer_heading_variations) is True
    assert parser._is_text_match("Supplier Information", parser.manufacturer_heading_variations) is True
    assert parser._is_text_match("Random Text", parser.manufacturer_heading_variations) is False


def test_extract_product_types():
    """Test extracting product types from a document."""
    # Create a test document with product types and manufacturers
    content_structure = [
        ("Introduction", "Heading 1"),
        ("Products", "Heading 1"),
        ("Product Type 1", "Heading 2"),
        ("Manufacturer: Company A", "Heading 3"),
        ("Description 1", "Heading 4"),
        ("Description 2", "Heading 4"),
        ("Product Type 2", "Heading 2"),
        ("Manufacturers: Company B, Company C", "Heading 3"),
        ("Description 3", "Heading 4"),
        ("Another Section", "Heading 1"),
    ]
    
    file_path = create_test_document("test_product_types.docx", content_structure)
    
    # Parse the document
    parser = DocumentParser()
    doc = Document(file_path)
    products_section = {"index": 1, "text": "Products"}
    product_types = parser._extract_product_types(doc, products_section)
    
    # Check that product types were extracted correctly
    assert len(product_types) == 2
    
    # Check first product type
    assert product_types[0]["name"] == "Product Type 1"
    assert len(product_types[0]["manufacturers"]) == 1
    assert product_types[0]["manufacturers"][0]["name"] == "Manufacturer: Company A"
    assert len(product_types[0]["manufacturers"][0]["descriptions"]) == 2
    
    # Check second product type
    assert product_types[1]["name"] == "Product Type 2"
    assert len(product_types[1]["manufacturers"]) == 1
    assert product_types[1]["manufacturers"][0]["name"] == "Manufacturers: Company B, Company C"
    assert len(product_types[1]["manufacturers"][0]["descriptions"]) == 1

def test_extract_product_types_with_style_variations():
    """Test extracting product types with different style variations."""
    # Create a test document with product types and manufacturers using different styles
    content_structure = [
        ("Introduction", "Heading 1"),
        ("Products", "Heading 1"),
        ("Product Type 1", "Title 2"),  # Using Title 2 instead of Heading 2
        ("Supplier: Company A", "H3"),  # Using H3 instead of Heading 3
        ("Description 1", "Title 4"),   # Using Title 4 instead of Heading 4
        ("Description 2", "Title 4"),
        ("Product Type 2", "Title 2"),
        ("Vendor: Company B", "H3"),
        ("Description 3", "Title 4"),
        ("Another Section", "Heading 1"),
    ]
    
    # Add custom styles to the test document
    file_path = create_test_document("test_style_variations.docx", content_structure)
    
    # Parse the document
    parser = DocumentParser()
    doc = Document(file_path)
    products_section = {"index": 1, "text": "Products"}
    product_types = parser._extract_product_types(doc, products_section)
    
    # Check that product types were extracted correctly despite style variations
    assert len(product_types) == 2
    
    # Check first product type
    assert product_types[0]["name"] == "Product Type 1"
    assert len(product_types[0]["manufacturers"]) == 1
    assert product_types[0]["manufacturers"][0]["name"] == "Supplier: Company A"
    assert len(product_types[0]["manufacturers"][0]["descriptions"]) == 2
    
    # Check second product type
    assert product_types[1]["name"] == "Product Type 2"
    assert len(product_types[1]["manufacturers"]) == 1
    assert product_types[1]["manufacturers"][0]["name"] == "Vendor: Company B"
    assert len(product_types[1]["manufacturers"][0]["descriptions"]) == 1


def test_parse_document_with_valid_structure():
    """Test parsing a document with a valid structure."""
    # Create a test document with a valid structure
    content_structure = [
        ("Introduction", "Heading 1"),
        ("This is an introduction", "Normal"),
        ("Products", "Heading 1"),
        ("Product Type 1", "Heading 2"),
        ("Manufacturer: Company A", "Heading 3"),
        ("Description 1", "Heading 4"),
        ("Product Type 2", "Heading 2"),
        ("Manufacturers: Company B, Company C", "Heading 3"),
        ("Description 2", "Heading 4"),
    ]
    
    file_path = create_test_document("test_valid_document.docx", content_structure)
    
    # Parse the document
    document_data = parse_document(file_path)
    
    # Check that the document was parsed correctly
    assert document_data is not None
    assert document_data["filename"] == "test_valid_document"
    assert document_data["products_section"] is not None
    assert len(document_data["product_types"]) == 2


def test_parse_document_with_missing_products_section():
    """Test parsing a document with a missing Products section."""
    # Create a test document without a Products section
    content_structure = [
        ("Introduction", "Heading 1"),
        ("This is an introduction", "Normal"),
        ("Section 1", "Heading 1"),
        ("Subsection 1", "Heading 2"),
        ("Details", "Normal"),
    ]
    
    file_path = create_test_document("test_missing_products.docx", content_structure)
    
    # Parse the document
    document_data = parse_document(file_path)
    
    # Check that the document was parsed but Products section is missing
    assert document_data is not None
    assert document_data["filename"] == "test_missing_products"
    assert document_data["products_section"] is None
    assert len(document_data["product_types"]) == 0


def test_parse_document_with_empty_product_types():
    """Test parsing a document with no product types."""
    # Create a test document with a Products section but no product types
    content_structure = [
        ("Introduction", "Heading 1"),
        ("This is an introduction", "Normal"),
        ("Products", "Heading 1"),
        ("Details about products", "Normal"),
        ("Conclusion", "Heading 1"),
    ]
    
    file_path = create_test_document("test_empty_products.docx", content_structure)
    
    # Parse the document
    document_data = parse_document(file_path)
    
    # Check that the document was parsed but no product types were found
    assert document_data is not None
    assert document_data["filename"] == "test_empty_products"
    assert document_data["products_section"] is not None
    assert len(document_data["product_types"]) == 0


def test_get_document_structure():
    """Test getting a summary of the document's structure."""
    # Create a test document
    content_structure = [
        ("Heading 1", "Heading 1"),
        ("Normal text", "Normal"),
        ("Heading 2", "Heading 2"),
        ("More normal text", "Normal"),
        ("Heading 3", "Heading 3"),
        ("Heading 4", "Heading 4"),
    ]
    
    file_path = create_test_document("test_structure.docx", content_structure)
    
    # Get document structure
    parser = DocumentParser()
    structure = parser.get_document_structure(file_path)
    
    # Check that the structure was analyzed correctly
    assert structure is not None
    assert structure["filename"] == "test_structure.docx"
    assert structure["heading_counts"]["Heading 1"] == 1
    assert structure["heading_counts"]["Heading 2"] == 1
    assert structure["heading_counts"]["Heading 3"] == 1
    assert structure["heading_counts"]["Heading 4"] == 1
    assert structure["paragraph_count"] == 6


if __name__ == "__main__":
    pytest.main(["-v", __file__])
